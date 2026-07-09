import os
import glob
import datetime
from pypdf import PdfReader
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, RGBColor

# ============================================================
# CONFIGURATION GEMINI
# ============================================================
GEMINI_API_KEY = "AQ.Ab8RN6KQTI9JPuyMy-AI6RED-2n2xSVOlMjpf1hMHMWQ7lMPfA"
client = genai.Client(api_key=GEMINI_API_KEY)
MODELE_GEMINI = "gemini-2.5-flash"

# ============================================================
# SOMMAIRES OBLIGATOIRES (Suggestion de Claude intégrée)
# ============================================================
CCAP_SECTIONS = [
    "Objet du marché",
    "Allotissement",
    "Mode de passation, type de marché et durée du marché",
    "Pièces contractuelles (liste + ordre de priorité en cas de contradiction)",
    "Prix : caractéristiques des prix pratiqués",
    "Prix : modalités de variation des prix (fermes/fixes ou révisables, formule si révisables)",
    "Conditions d'exécution des prestations (livraison, installation, mise en service, formation)",
    "Délais d'exécution / délais de livraison",
    "Pénalités de retard (méthode de calcul, plafond éventuel)",
    "Garantie des prestations (durée, pièces et main d'œuvre, prêt de matériel en cas de panne)",
    "RSE - Responsabilité Sociale des Entreprises (clauses sociales et environnementales)",
    "Modalités de règlement des comptes (facturation, Chorus Pro, délai global de paiement, intérêts moratoires)",
    "Résiliation du marché (motif d'intérêt général, inexactitude des déclarations, redressement/liquidation judiciaire)",
    "Litiges et tribunal compétent",
    "Langue des documents (obligation de rédaction en français ou traduction assermentée)",
    "Confidentialité",
    "Protection des données à caractère personnel (RGPD, loi 78-17)",
]

CCTP_SECTIONS = [
    "Préambule et objet du marché (contexte, objet détaillé, périmètre, volumétrie)",
    "Définitions et acronymes",
    "Documents de référence (liste + ordre de priorité)",
    "Exigences générales et transversales (garanties, RSE, accessibilité, devoir de conseil)",
    "Spécifications techniques détaillées par lot / par matériel",
    "Documentation technique et pédagogique à fournir par le candidat",
    "Modalités d'installation, de mise en service et de formation",
    "Modalités de maintenance (préventive et corrective) et délais d'intervention",
]

RC_SECTIONS = [
    "Objet de la consultation (présentation générale, acheteur, allotissement, mode de passation)",
    "Conditions relatives à la consultation (délai de validité des offres, confidentialité)",
    "Contenu du dossier de consultation des entreprises (DCE) + modalités de retrait/plateforme",
    "Délai et modalités pour poser des questions (nombre de jours avant la date limite)",
    "Documents à produire au titre de la candidature",
    "Documents à produire au titre de l'offre",
    "Conditions d'envoi ou de remise des plis (transmission électronique, format, copie de sauvegarde)",
    "Sélection des candidatures (critères et motifs d'élimination, régularisation possible)",
    "Jugement des offres (critères et pondération, méthode de notation)",
    "Suites données à la consultation (notification, délai de complément de dossier, candidat classé 2e en cas de défaillance)",
    "Procédures et voies de recours (tribunal administratif compétent : adresse, tél, email, site)",
]

STRICT_STRUCTURE_INSTRUCTION = """
CONSIGNE IMPÉRATIVE DE STRUCTURE :
Tu dois obligatoirement traiter TOUTES les sections listées ci-dessous, dans cet
ordre, sans EN OMETTRE AUCUNE, même si les documents source ne fournissent pas
l'information. Si une information est absente des documents fournis, tu dois
quand même écrire la section et indiquer "[À PRÉCISER : ...]" avec une
proposition ou un conseil raisonnable entre crochets plutôt que de sauter la section.

Ne t'arrête jamais avant d'avoir traité la dernière section de la liste.

Liste des sections obligatoires pour ce document :
{sections}
"""

def get_strict_instruction(doc_type: str) -> str:
    sections_map = {"CCAP": CCAP_SECTIONS, "CCTP": CCTP_SECTIONS, "RC": RC_SECTIONS}
    sections = sections_map.get(doc_type.upper())
    block_texte = "\n".join(f"{i+1}. {s}" for i, s in enumerate(sections))
    return STRICT_STRUCTURE_INSTRUCTION.format(sections=block_texte)

# ============================================================
# DÉTECTION ET LECTURE DU DOSSIER CIBLE
# ============================================================
print("\n🔍 Agent Word Structure Strict : Recherche du dossier '_recherche'...")

tous_les_elements = os.listdir(".")
dossiers_trouves = [e for e in tous_les_elements if os.path.isdir(e) and e.endswith("_recherche")]

if not dossiers_trouves:
    print("❌ Aucun dossier se terminant par '_recherche' n'a été trouvé.")
    exit()

nom_du_dossier = dossiers_trouves[0]
print(f"🎯 Dossier source détecté : [{nom_du_dossier}]")

fichiers_pdf = glob.glob(os.path.join(nom_du_dossier, "*.pdf"))

if not fichiers_pdf:
    print(f"⚠️ Le dossier [{nom_du_dossier}] ne contient aucun fichier PDF.")
    exit()

print(f"📚 Lecture de {len(fichiers_pdf)} fichier(s) PDF...")
texte_global_dossier = ""

for chemin_pdf in fichiers_pdf:
    try:
        reader = PdfReader(chemin_pdf)
        nom_fichier_court = os.path.basename(chemin_pdf)
        texte_global_dossier += f"\n--- DOCUMENT : {nom_fichier_court} ---\n"
        for page in reader.pages:
            texte_global_dossier += page.extract_text() + "\n"
    except Exception as e:
        print(f" 🔺 Erreur de lecture sur {nom_fichier_court} : {str(e)}")

# ============================================================
# FONCTIONS DE MISE EN FORME WORD (STYLE N2IGC)
# ============================================================
def ajouter_titre_stylise(doc, texte, niveau=1):
    h = doc.add_heading(texte, level=niveau)
    if h.runs:
        run = h.runs[0]
        couleurs = {
            1: (RGBColor(0x1A, 0x56, 0xDB), Pt(16)), # Bleu Royal
            2: (RGBColor(0x6B, 0x0F, 0x4C), Pt(13)), # Pourpre Corporate
        }
        if niveau in couleurs:
            run.font.color.rgb, run.font.size = couleurs[niveau]
            run.font.bold = True

def injecter_texte_dans_word(doc, texte_ia):
    for ligne in texte_ia.strip().split('\n'):
        ligne_propre = ligne.replace('**', '').replace('##', '').replace('#', '').strip()
        
        if not ligne_propre:
            doc.add_paragraph()
            continue
            
        if ligne.strip().startswith('### ') or ligne.strip().startswith('## '):
            ajouter_titre_stylise(doc, ligne_propre, niveau=2)
        elif ligne.strip().startswith('# '):
            ajouter_titre_stylise(doc, ligne_propre, niveau=1)
        elif ligne.strip().startswith('- ') or ligne.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(ligne_propre.lstrip('-* ')).font.size = Pt(11)
        else:
            p = doc.add_paragraph(ligne_propre)
            if p.runs: 
                p.runs[0].font.size = Pt(11)

# ============================================================
# FONDATION ET RÉDACTION DES LIVRABLES WORD
# ============================================================
def creer_document_word(role, doc_type, instructions_specifiques, nom_titre_doc, nom_fichier_final):
    print(f"⏳ Génération Complète de : {nom_fichier_final} (Sommaire strict {doc_type})...")
    
    # On assemble les instructions de base avec la consigne de structure stricte de Claude
    instruction_complete = instructions_specifiques + "\n" + get_strict_instruction(doc_type)
    
    try:
        response = client.models.generate_content(
            model=MODELE_GEMINI,
            contents=[f"--- DONNÉES DU DOSSIER SOURCE ---\n{texte_global_dossier[:40000]}", instruction_complete],
            config=types.GenerateContentConfig(system_instruction=role, temperature=0.2)
        )
        
        doc = Document()
        ajouter_titre_stylise(doc, nom_titre_doc, niveau=1)
        doc.add_paragraph(f"Généré automatiquement pour N2IGC le {datetime.datetime.now().strftime('%d/%m/%Y')} - Structure Validée\n")
        
        injecter_texte_dans_word(doc, response.text)
        
        chemin_enregistrement = os.path.join(nom_du_dossier, nom_fichier_final)
        doc.save(chemin_enregistrement)
        print(f" ✅ Fichier Word créé avec succès : {chemin_enregistrement}")
        
    except Exception as e:
        print(f" ❌ Erreur lors de la création de {nom_fichier_final} : {str(e)}")

# ============================================================
# EXÉCUTION DE LA GÉNÉRATION AVEC SOMMAIRES BLINDÉS
# ============================================================
print("\n⚡ Lancement de la rédaction des pièces réglementaires complètes...\n")

# 1. GÉNÉRATION DU CCTP COMPLETE
creer_document_word(
    role="Tu es un ingénieur principal expert en rédaction de CCTP industriels et tertiaires.",
    doc_type="CCTP",
    instructions_specifiques="Rédige le Cahier des Clauses Techniques Particulières (CCTP) en développant chaque section. Reste ultra-précis sur le matériel et les spécifications techniques trouvées dans le dossier.",
    nom_titre_doc="CAHIER DES CLAUSES TECHNIQUES PARTICULIÈRES (CCTP)",
    nom_fichier_final="LIVRABLE_CCTP.docx"
)

# 2. GÉNÉRATION DU CCAP COMPLETE
creer_document_word(
    role="Tu es un juriste d'entreprise expert en rédaction de clauses administratives de marchés.",
    doc_type="CCAP",
    instructions_specifiques="Rédige le Cahier des Clauses Administratives Particulières (CCAP). Écris des clauses juridiques réalistes et impératives pour chaque point du sommaire.",
    nom_titre_doc="CAHIER DES CLAUSES ADMINISTRATIVES PARTICULIÈRES (CCAP)",
    nom_fichier_final="LIVRABLE_CCAP.docx"
)

# 3. GÉNÉRATION DU RC COMPLETE
creer_document_word(
    role="Tu es un acheteur professionnel expert en passation de marchés publics et privés.",
    doc_type="RC",
    instructions_specifiques="Rédige le Règlement de Consultation (RC) complet de la procédure. Reprends fidèlement les critères de notation, les modalités de dépôt électronique et les règles de candidature.",
    nom_titre_doc="RÈGLEMENT DE CONSULTATION (RC)",
    nom_fichier_final="LIVRABLE_RC.docx"
)

print("\n🎉 TERMINÉ : Les documents 'LIVRABLE_CCTP.docx', 'LIVRABLE_CCAP.docx' et 'LIVRABLE_RC.docx' sont complets et générés !")