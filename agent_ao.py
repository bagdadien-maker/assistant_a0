import os
import time
from pypdf import PdfReader
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# ============================================================
# CONFIGURATION
# ============================================================
API_KEY = "AQ.Ab8RN6KQTI9JPuyMy-AI6RED-2n2xSVOlMjpf1hMHMWQ7lMPfA"
client = genai.Client(api_key=API_KEY)

MODELES = [
    "gemini-2.5-flash",
    "gemini-1.5-flash",
    "gemini-1.5-pro",
]

# ============================================================
# FONCTIONS UTILITAIRES
# ============================================================

def extraire_texte_pdf(chemin_pdf):
    if not os.path.exists(chemin_pdf):
        print(f"   ⚠️  Fichier introuvable : {chemin_pdf}")
        return ""
    reader = PdfReader(chemin_pdf)
    texte = ""
    for page in reader.pages:
        texte += page.extract_text() + "\n"
    print(f"   ✅ {chemin_pdf} lu ({len(reader.pages)} pages)")
    return texte


def appeler_ia(prompt, system="Tu es un expert en marchés publics français.", label=""):
    """Appelle Gemini 2.0 Flash en priorité avec fallback automatique."""
    for modele in MODELES:
        for tentative in range(3):
            try:
                response = client.models.generate_content(
                    model=modele,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=system,
                        temperature=0.2,
                    )
                )
                texte = response.text if response.text else ""

                # Détection réponse vide ou corrompue (suite de 0 ou trop courte)
                if len(texte.strip()) < 100 or texte.count("0") > len(texte) * 0.3:
                    print(f"   ⚠️  Réponse invalide détectée ({modele}) — relance...")
                    time.sleep(10)
                    continue  # relance le même modèle

                if modele != MODELES[0]:
                    print(f"   ℹ️  Modèle de secours utilisé : {modele}")
                return texte

            except Exception as e:
                err = str(e)
                if "429" in err or "quota" in err.lower():
                    if tentative < 2:
                        attente = (tentative + 1) * 15
                        print(f"   ⏳ [{modele}] surchargé — attente {attente}s (essai {tentative+1}/3)...")
                        time.sleep(attente)
                    else:
                        print(f"   🔄 [{modele}] abandon — modèle suivant...")
                        break
                else:
                    print(f"\n   ❌ Erreur critique ({label}) : {err}")
                    raise

    raise Exception("❌ Tous les modèles sont indisponibles. Réessaie dans quelques minutes.")


def prompt_produits_lot(nom_lot, description_lot, cctp_extrait):
    """Génère le prompt de recherche produits pour UN seul lot."""
    return f"""
Tu es un commercial expert en marchés publics et équipements industriels/pédagogiques.

Analyse les exigences du lot suivant et propose des produits réels disponibles en France ou en Europe.
Comme les fabricants ne mettent pas toujours leurs fiches précises ou leurs prix en ligne, donne les informations les plus fidèles au marché actuel.

=== LOT À TRAITER ===
{nom_lot}

=== EXIGENCES TECHNIQUES ===
{description_lot}

=== CONTEXTE CCTP ===
{cctp_extrait[:15000]}

INSTRUCTIONS — Pour ce lot, tu dois absolument fournir :

## Produit recommandé n°1
- **Marque et référence** : [Marque réelle + référence exacte ou catalogue du fabricant]
- **Description** : [2-3 lignes sur le matériel]
- **Conformité CCTP** : [Vérification point par point avec les exigences du texte]
- **Prix estimé HT** : [Fourchette de prix réaliste en € basée sur le marché]
- **Site du fabricant** : [URL globale ou principale du site web officiel du constructeur]
- **Disponibilité doc technique** : [Mettre le lien direct de la page produit si elle existe, sinon écrire "Sur demande de devis / Catalogue commercial réservé aux enseignants"]
- **Action de sourcing** : [Conseil court sur la démarche à faire pour obtenir le prix exact, ex: "Contacter le technico-commercial régional de chez X"]

## Produit recommandé n°2
[même structure]

## Produit recommandé n°3 (optionnel)
[même structure]

## Meilleur choix recommandé
[Lequel choisir parmi les propositions et pourquoi en 3-4 lignes]

## Estimation de prix pour ce lot
- Prix matériel HT : [montant ou fourchette]
- Installation + livraison : [montant estimé]
- Formation : [montant estimé]
- **TOTAL LOT HT** : [montant total]

Sois précis et réaliste. Ne mets jamais de zéros ou de texte vide.
"""


def ajouter_titre(doc, texte, niveau=1):
    h = doc.add_heading(texte, level=niveau)
    if h.runs:
        run = h.runs[0]
        if niveau == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
            run.font.size = Pt(14)
        elif niveau == 2:
            run.font.color.rgb = RGBColor(0x0E, 0x9F, 0x6E)
            run.font.size = Pt(12)
        elif niveau == 3:
            run.font.color.rgb = RGBColor(0xE3, 0x7A, 0x00)
            run.font.size = Pt(11)


def ajouter_separateur(doc):
    doc.add_paragraph("─" * 80)


def ecrire_lignes(doc, texte):
    for ligne in texte.strip().split('\n'):
        propre = ligne.replace('**', '').replace('##', '').replace('#', '').strip()
        if not propre:
            doc.add_paragraph()
            continue
        if ligne.strip().startswith('## ') or ligne.strip().startswith('# '):
            ajouter_titre(doc, propre, niveau=2)
        elif ligne.strip().startswith('### '):
            ajouter_titre(doc, propre, niveau=3)
        elif ligne.strip().startswith('- ') or ligne.strip().startswith('* '):
            contenu = propre.lstrip('-* ')
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(contenu)
            run.font.size = Pt(11)
            if 'http' in contenu:
                run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
            elif '€' in contenu or 'prix' in contenu.lower() or 'total' in contenu.lower():
                run.font.color.rgb = RGBColor(0x0E, 0x9F, 0x6E)
                run.bold = True
        else:
            p = doc.add_paragraph(propre)
            if p.runs:
                p.runs[0].font.size = Pt(11)


# ============================================================
# DÉMARRAGE
# ============================================================
print("\n" + "="*60)
print("   AGENT IA — ANALYSE APPEL D'OFFRES")
print(f"   Modèle principal : {MODELES[0]}")
print(f"   Fallback : {' → '.join(MODELES[1:])}")
print("="*60)

# ============================================================
# LECTURE DES PDFs
# ============================================================
print("\n📖 Lecture des fichiers PDF...\n")
texte_rc   = extraire_texte_pdf("RC.pdf")
texte_ccap = extraire_texte_pdf("CCAP.pdf")
texte_cctp = extraire_texte_pdf("CCTP.pdf")

if not texte_cctp:
    print("\n❌ CCTP.pdf introuvable. Place RC.pdf, CCAP.pdf et CCTP.pdf dans le même dossier.")
    exit()

# ============================================================
# ÉTAPE 1 — FILTRAGE
# ============================================================
print("\n🔍 Étape 1/5 — Filtrage éligibilité...")

prompt_filtrage = f"""
Analyse ces documents et réponds avec ce format EXACT :

DÉCISION: [ÉLIGIBLE ou REJETÉ]
RAISON: [explication courte en 2 phrases]
MONTANT_ESTIMÉ: [montant trouvé ou "Non précisé"]
VISITE_OBLIGATOIRE: [Oui / Non / Non précisé]
DATE_LIMITE_OFFRES: [date ou "Non précisée"]
DURÉE_MARCHÉ: [durée ou "Non précisée"]
CRITÈRES_JUGEMENT:
- Critère : [nom] — Pondération : [%]

--- RC ---
{texte_rc[:12000]}
--- CCAP ---
{texte_ccap[:12000]}
"""
resultat_filtrage = appeler_ia(prompt_filtrage, label="Filtrage")
print("   ✅ Filtrage terminé")

# ============================================================
# ÉTAPE 2 — ANALYSE TECHNIQUE
# ============================================================
print("\n🔧 Étape 2/5 — Analyse technique du CCTP...")

prompt_technique = f"""
Analyse ce CCTP et rédige un cahier des charges ultra-détaillé par lot/équipement.
Pour CHAQUE lot identifié, structure ainsi :

## [Nom du lot / équipement]
- Quantité : [nombre]
- Contraintes fonctionnelles : [liste complète]
- Contraintes technologiques (specs exactes) : [liste complète]
- Documents à fournir : [liste]
- Points de vigilance : [liste]

Reprends TOUTES les contraintes sans exception.

--- CCTP ---
{texte_cctp[:25000]}
"""
resultat_technique = appeler_ia(prompt_technique, label="Technique")
print("   ✅ Analyse technique terminée")

# ============================================================
# ÉTAPE 3A — PRODUITS LOT 1 (banc moteur B6.7 + oscilloscope)
# ============================================================
print("\n🛒 Étape 3a/5 — Produits Lot 1 (Banc B6.7 + oscilloscope)...")

desc_lot1 = """
LOT 1 : Banc moteur didactique type Cummins B6.7 Stage V + oscilloscope 4 voies
- Moteur diesel Cummins B6.7 ou équivalent, Stage V
- Injection Common Rail, turbo géométrie variable
- Post-traitement DOC + DPF + SCR + AdBlue
- Boîte à pannes (circuits ouverts, courts-circuits, résistances)
- Panneau de mesures (résistances, tensions, signaux capteurs)
- Pupitre commande complet (démarrage sécurisé, arrêt urgence, CAN/J1939, accélération pied/main)
- Bac de rétention, protections machines 2006/42/CE
- Oscilloscope 4 voies numérique Windows 11, USB 1.5m min
- Accessoires : pinces ampèremétriques, cordons banane Ø4mm, sondes x10, cordons fusibles
- Logiciel avec tests guidés et base de données diagnostic
- Dossier pédagogique TP élève + enseignant inclus
- Documentation technique française sur clé USB
"""

prompt_lot1 = prompt_produits_lot(
    "LOT 1 — Banc moteur B6.7 Stage V + Oscilloscope 4 voies",
    desc_lot1,
    texte_cctp
)
resultat_produits_lot1 = appeler_ia(
    prompt_lot1,
    system="Tu es un expert commercial en équipements didactiques et pédagogiques pour lycées agricoles et techniques. Tu connais parfaitement les fournisseurs français : DIDAC BDH, Exxotest, ConsuLab, Pico Technology, Jaltest, Haldex, etc. Tu donnes des URLs réelles, des prix réalistes et des références exactes.",
    label="Produits Lot 1"
)
print("   ✅ Produits Lot 1 terminés")

# Pause entre les deux appels pour éviter la surcharge
time.sleep(5)

# ============================================================
# ÉTAPE 3B — PRODUITS LOT 2 (banc moteur F3.8 + acquisition 16 voies)
# ============================================================
print("\n🛒 Étape 3b/5 — Produits Lot 2 (Banc F3.8 + acquisition 16 voies)...")

desc_lot2 = """
LOT 2 : Banc moteur didactique type Cummins B3.8 Stage V + système acquisition 16 voies
- Moteur diesel Cummins B3.8 ou équivalent, Stage V
- Injection Common Rail, turbo géométrie fixe
- Post-traitement DOC + DPF + SCR + AdBlue
- Boîte à pannes (circuits ouverts, courts-circuits, résistances)
- Panneau de mesures (résistances, tensions, signaux capteurs)
- Pupitre commande complet (démarrage sécurisé, arrêt urgence, CAN/J1939, accélération pied/main)
- Bac de rétention, protections machines 2006/42/CE
- Système acquisition données : 16 voies analogiques simultanées + 8 voies logiques/différentielles
- Fréquence acquisition minimum 100 kHz
- Alimentation USB (pas de 230V dédié)
- Protections courts-circuits, surcharges, inversions de branchement
- Connexions fiches bananes Ø4mm
- Fonction oscilloscope multivoies 16 voies simultanées
- Logiciel avec aide câblage, déclenchement, graphiques, calculs
- Dossier pédagogique TP élève + enseignant inclus
- Documentation technique française sur clé USB
"""

prompt_lot2 = prompt_produits_lot(
    "LOT 2 — Banc moteur B3.8 Stage V + Système acquisition données 16 voies",
    desc_lot2,
    texte_cctp
)
resultat_produits_lot2 = appeler_ia(
    prompt_lot2,
    system="Tu es un expert commercial en équipements didactiques et pédagogiques pour lycées agricoles et techniques. Tu connais parfaitement les fournisseurs français : DIDAC BDH, Exxotest, ConsuLab, National Instruments, Digilent, Dewesoft, Sefram, etc. Tu donnes des URLs réelles, des prix réalistes et des références exactes.",
    label="Produits Lot 2"
)
print("   ✅ Produits Lot 2 terminés")

# ============================================================
# ÉTAPE 4 — TABLEAU RÉCAPITULATIF GLOBAL
# ============================================================
print("\n📊 Étape 4/5 — Tableau récapitulatif global et total estimé...")

time.sleep(3)

prompt_recap = f"""
Sur la base de ces deux analyses produits, rédige un tableau récapitulatif global.

--- PRODUITS LOT 1 ---
{resultat_produits_lot1[:5000]}

--- PRODUITS LOT 2 ---
{resultat_produits_lot2[:5000]}

Rédige :

## TABLEAU RÉCAPITULATIF GLOBAL

| Lot | Équipement | Produit recommandé | Fournisseur | Prix HT estimé |
|-----|-----------|-------------------|-------------|----------------|
| Lot 1 | Banc moteur B6.7 | ... | ... | ... € |
| Lot 1 | Oscilloscope 4 voies | ... | ... | ... € |
| Lot 2 | Banc moteur B3.8 | ... | ... | ... € |
| Lot 2 | Acquisition 16 voies | ... | ... | ... € |
| - | Formation professeurs | ... | ... | ... € |

**TOTAL ESTIMÉ HT : ... €**
**TVA 20% : ... €**
**TOTAL TTC : ... €**

## Positionnement tarifaire recommandé
[Conseil sur le prix à proposer pour être compétitif sur ce marché public]
"""

resultat_recap = appeler_ia(prompt_recap, label="Récapitulatif")
print("   ✅ Récapitulatif terminé")

# ============================================================
# ÉTAPE 5 — SYNTHÈSE STRATÉGIQUE
# ============================================================
print("\n💡 Étape 5/5 — Synthèse stratégique + rédaction Word...")

time.sleep(3)

prompt_synthese = f"""
Rédige une synthèse stratégique (300 mots max) :

## Faut-il répondre ?
[recommandation + justification]

## Les 3 points clés pour gagner
- Point 1
- Point 2  
- Point 3

## Budget total recommandé
[montant HT compétitif à proposer]

## Risques principaux
- Risque 1
- Risque 2

Contexte :
{resultat_filtrage[:1000]}
{resultat_recap[:1000]}
"""
resultat_synthese = appeler_ia(prompt_synthese, label="Synthèse")

# ============================================================
# RÉDACTION DU RAPPORT WORD
# ============================================================
doc = Document()

# Page de garde
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("RAPPORT D'ANALYSE — APPEL D'OFFRES")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_paragraph()
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run(
    "Produits recommandés · Fournisseurs · Fiches techniques · Prix estimés\n"
    f"Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}"
)
doc.add_page_break()

# Section 1 : Filtrage
ajouter_titre(doc, "1. FILTRAGE — ÉLIGIBILITÉ DU MARCHÉ", niveau=1)
doc.add_paragraph()
for ligne in resultat_filtrage.strip().split('\n'):
    propre = ligne.strip()
    if not propre:
        continue
    if propre.startswith("DÉCISION:"):
        decision = propre.replace("DÉCISION:", "").strip()
        p = doc.add_paragraph()
        run = p.add_run(f"➤ DÉCISION : {decision}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = (
            RGBColor(0x0E, 0x9F, 0x6E) if "ÉLIGIBLE" in decision
            else RGBColor(0xE0, 0x2A, 0x2A)
        )
    elif ":" in propre and not propre.startswith("-"):
        cle, _, val = propre.partition(":")
        p = doc.add_paragraph()
        r1 = p.add_run(f"{cle.strip()} : ")
        r1.bold = True
        r1.font.size = Pt(11)
        p.add_run(val.strip()).font.size = Pt(11)
    elif propre.startswith("-"):
        doc.add_paragraph(propre.lstrip("- "), style="List Bullet")
    else:
        doc.add_paragraph(propre)

doc.add_page_break()

# Section 2 : Cahier des charges
ajouter_titre(doc, "2. CAHIER DES CHARGES TECHNIQUE DÉTAILLÉ", niveau=1)
doc.add_paragraph()
ecrire_lignes(doc, resultat_technique)
doc.add_page_break()

# Section 3 : Produits Lot 1
ajouter_titre(doc, "3. PRODUITS RECOMMANDÉS — LOT 1 (Banc B6.7 + Oscilloscope)", niveau=1)
doc.add_paragraph()
p_intro1 = doc.add_paragraph()
r_intro1 = p_intro1.add_run(
    "Produits réels disponibles en France pour le banc moteur Cummins B6.7 Stage V "
    "et l'oscilloscope 4 voies, avec fournisseurs, fiches techniques et prix estimés."
)
r_intro1.font.size = Pt(11)
r_intro1.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
r_intro1.italic = True
doc.add_paragraph()
ecrire_lignes(doc, resultat_produits_lot1)
doc.add_page_break()

# Section 4 : Produits Lot 2
ajouter_titre(doc, "4. PRODUITS RECOMMANDÉS — LOT 2 (Banc B3.8 + Acquisition 16 voies)", niveau=1)
doc.add_paragraph()
p_intro2 = doc.add_paragraph()
r_intro2 = p_intro2.add_run(
    "Produits réels disponibles en France pour le banc moteur Cummins B3.8 Stage V "
    "et le système d'acquisition 16 voies, avec fournisseurs, fiches techniques et prix estimés."
)
r_intro2.font.size = Pt(11)
r_intro2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
r_intro2.italic = True
doc.add_paragraph()
ecrire_lignes(doc, resultat_produits_lot2)
doc.add_page_break()

# Section 5 : Récapitulatif global
ajouter_titre(doc, "5. TABLEAU RÉCAPITULATIF GLOBAL & BUDGET TOTAL", niveau=1)
doc.add_paragraph()
ecrire_lignes(doc, resultat_recap)
doc.add_page_break()

# Section 6 : Synthèse
ajouter_titre(doc, "6. SYNTHÈSE STRATÉGIQUE", niveau=1)
doc.add_paragraph()
ecrire_lignes(doc, resultat_synthese)

# Pied de page
doc.add_paragraph()
ajouter_separateur(doc)
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run(
    f"Document généré par Agent IA — Analyse Marchés Publics · "
    f"{datetime.datetime.now().strftime('%d/%m/%Y')}"
)
rf.font.size = Pt(9)
rf.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

# Sauvegarde
nom_docx = f"Rapport_AO_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
doc.save(nom_docx)

with open("Rapport_AO_complet.txt", "w", encoding="utf-8") as f:
    f.write("=== FILTRAGE ===\n\n" + resultat_filtrage + "\n\n")
    f.write("=== ANALYSE TECHNIQUE ===\n\n" + resultat_technique + "\n\n")
    f.write("=== PRODUITS LOT 1 ===\n\n" + resultat_produits_lot1 + "\n\n")
    f.write("=== PRODUITS LOT 2 ===\n\n" + resultat_produits_lot2 + "\n\n")
    f.write("=== RÉCAPITULATIF ===\n\n" + resultat_recap + "\n\n")
    f.write("=== SYNTHÈSE ===\n\n" + resultat_synthese)

print(f"\n{'='*60}")
print(f"   ✅ RAPPORT GÉNÉRÉ AVEC SUCCÈS !")
print(f"   📄 Word  : {nom_docx}")
print(f"   📄 Texte : Rapport_AO_complet.txt")
print(f"{'='*60}\n")