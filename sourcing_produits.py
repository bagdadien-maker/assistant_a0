import os
import datetime
from pypdf import PdfReader
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, RGBColor

# ============================================================
# CONFIGURATION GEMINI
# ============================================================
GEMINI_API_KEY = "AQ.Ab8RN6KQTI9JPuyMy-AI6RED-2n2xSVOlMjpf1hMHMWQ7lMPfA"  # Mets ta clé ici
client = genai.Client(api_key=GEMINI_API_KEY)
MODELE_GEMINI = "gemini-2.5-flash"

def extraire_texte_pdf(chemin_pdf):
    if not os.path.exists(chemin_pdf): return ""
    reader = PdfReader(chemin_pdf)
    return "".join([page.extract_text() + "\n" for page in reader.pages])

def appeler_gemini(prompt, system_instruction):
    try:
        response = client.models.generate_content(
            model=MODELE_GEMINI,
            contents=prompt,
            config=types.GenerateContentConfig(system_instruction=system_instruction, temperature=0.2)
        )
        return response.text if response.text else ""
    except Exception as e:
        print(f"   ❌ Erreur Gemini : {str(e)}")
        return ""

def ajouter_titre(doc, texte, niveau=1):
    h = doc.add_heading(texte, level=niveau)
    if h.runs:
        run = h.runs[0]
        run.font.color.rgb = RGBColor(0x0E, 0x9F, 0x6E) if niveau==1 else RGBColor(0x1A, 0x56, 0xDB)
        run.font.size = Pt(14 if niveau==1 else 12)

def ecrire_lignes(doc, texte):
    for ligne in texte.strip().split('\n'):
        propre = ligne.replace('**', '').replace('##', '').replace('#', '').strip()
        if not propre: continue
        if ligne.strip().startswith('## ') or ligne.strip().startswith('# '):
            ajouter_titre(doc, propre, niveau=2)
        elif ligne.strip().startswith('- ') or ligne.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(propre.lstrip('-* '))
            run.font.size = Pt(11)
            if '€' in propre or 'total' in propre.lower():
                run.font.color.rgb = RGBColor(0x0E, 0x9F, 0x6E)
                run.bold = True
        else:
            p = doc.add_paragraph(propre)
            if p.runs: p.runs[0].font.size = Pt(11)

# ============================================================
# EXÉCUTION SOURCING
# ============================================================
print("\n🌐 DÉMARRAGE DU PROGRAMME 2 : SOURCING 3 PRODUITS MINIMUM (GEMINI)")
texte_cctp = extraire_texte_pdf("CCTP.pdf")

if not texte_cctp:
    print("❌ Le fichier CCTP.pdf est requis pour chercher les produits.")
    exit()

# Le prompt mis à jour pour exiger obligatoirement 3 produits par demande technique détectée
prompt_sourcing = f"""
Agis en tant que Directeur du Sourcing International pour N2IGC.
Analyse les besoins matériels et les demandes du CCTP fournis. 
Pour CHAQUE équipement ou lot majeur identifié, tu dois OBLIGATOIREMENT proposer au moins TROIS (3) produits réels différents, disponibles auprès de marques ou distributeurs en France/Europe (ex: Exxotest, DIDAC BDH, Pico Technology, Sefram, Cummins, National Instruments, Dewesoft, Jaltest, etc.).

Structure ta réponse scrupuleusement de cette manière pour chaque lot/demande :

## LOT [Numéro] : [Nom du Lot]

### 1️⃣ Produit Recommandé n°1 (Choix Principal)
- Marque et Référence exacte : 
- Description technique : [Comment il répond point par point au CCTP]
- Prix Estimé HT : [Fourchette en Euros, ex: 15 000€ - 18 000€]
- Site du fournisseur / Distributeur : [URL ou nom du site]

### 2️⃣ Produit Recommandé n°2 (Alternative Directe)
- Marque et Référence exacte : 
- Description technique : [Ses forces par rapport au besoin]
- Prix Estimé HT : [Fourchette en Euros]
- Site du fournisseur / Distributeur : [URL ou nom du site]

### 3️⃣ Produit Recommandé n°3 (Alternative Économique ou Haut de Gamme)
- Marque et Référence exacte : 
- Description technique : [Son positionnement]
- Prix Estimé HT : [Fourchette en Euros]
- Site du fournisseur / Distributeur : [URL ou nom du site]

### 🎯 Meilleur choix N2IGC
[Explique en 3 lignes lequel de ces 3 produits N2IGC doit intégrer dans son offre finale et pourquoi]

---

## TABLEAU RÉCAPITULATIF ET BUDGET ESTIMÉ
Génère un tableau Markdown propre listant : Lot | Équipement | Produit Retenu (Meilleur choix) | Fournisseur | Prix Estimé HT.
Ajoute à la fin les lignes :
- TOTAL MATÉRIEL HT
- ESTIMATION LIVRAISON & INSTALLATION HT
- TOTAL ESTIMÉ DE L'OFFRE HT

--- EXTRAIT CCTP ---
{texte_cctp[:15000]}
"""

print("\n🛒 Recherche, comparaison et sourcing des 3 options par équipement via Gemini...")
resultat_sourcing = appeler_gemini(prompt_sourcing, "Tu es un acheteur industriel B2B senior et un as du sourcing d'équipements didactiques, mécaniques et scientifiques.")

# CRÉATION DU RAPPORT DE SOURCING
print("\n📝 Génération du catalogue de sourcing de triple choix...")
doc = Document()
doc.add_heading("N2IGC - OFFRE COMMERCIALE COMPLÈTE (3 OPTIONS PAR BESOIN)", level=1)
doc.add_paragraph(f"Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}\n")

ajouter_titre(doc, "1. CATALOGUE COMPARATIF DE SOURCING (3 PRODUITS MINIMUM)", niveau=1)
ecrire_lignes(doc, resultat_sourcing)

doc.save("N2IGC_Sourcing_Produits.docx")
with open("N2IGC_Sourcing_Produits.txt", "w", encoding="utf-8") as f:
    f.write(resultat_sourcing)

print("✅ PROGRAMME 2 TERMINÉ : Fichiers 'N2IGC_Sourcing_Produits.docx' et '.txt' créés avec le triple choix !")