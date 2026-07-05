import os
import datetime
from pypdf import PdfReader
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# CONFIGURATION GEMINI
# ============================================================
GEMINI_API_KEY = "AQ.Ab8RN6KQTI9JPuyMy-AI6RED-2n2xSVOlMjpf1hMHMWQ7lMPfA"  # Mets ta clé ici
client = genai.Client(api_key=GEMINI_API_KEY)
MODELE_GEMINI = "gemini-2.5-flash"  # Modèle ultra-rapide et intelligent

# ============================================================
# FONCTIONS OUTILS
# ============================================================
def extraire_texte_pdf(chemin_pdf):
    if not os.path.exists(chemin_pdf):
        print(f"   ⚠️  Fichier introuvable : {chemin_pdf}")
        return ""
    reader = PdfReader(chemin_pdf)
    texte = ""
    for page in reader.pages:
        texte += page.extract_text() + "\n"
    print(f"   ✅ {chemin_pdf} lu ({len(reader.pages)} pages)")
    return texte

def appeler_gemini(prompt, system_instruction):
    try:
        response = client.models.generate_content(
            model=MODELE_GEMINI,
            contents=prompt,
            config=types.GenerateContentConfig(
                system_instruction=system_instruction,
                temperature=0.2,
            )
        )
        return response.text if response.text else "Aucune réponse générée."
    except Exception as e:
        print(f"   ❌ Erreur Gemini : {str(e)}")
        return ""

# ============================================================
# MISE EN FORME WORD
# ============================================================
def ajouter_titre(doc, texte, niveau=1):
    h = doc.add_heading(texte, level=niveau)
    if h.runs:
        run = h.runs[0]
        couleurs = {
            1: (RGBColor(0x1A, 0x56, 0xDB), Pt(14)),
            2: (RGBColor(0x6B, 0x0F, 0x4C), Pt(12)), # Pourpre N2IGC
        }
        if niveau in couleurs:
            run.font.color.rgb, run.font.size = couleurs[niveau]

def ecrire_lignes(doc, texte):
    for ligne in texte.strip().split('\n'):
        propre = ligne.replace('**', '').replace('##', '').replace('#', '').strip()
        if not propre:
            doc.add_paragraph()
            continue
        if ligne.strip().startswith('## ') or ligne.strip().startswith('# '):
            ajouter_titre(doc, propre, niveau=2)
        elif ligne.strip().startswith('- ') or ligne.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(propre.lstrip('-* ')).font.size = Pt(11)
        else:
            p = doc.add_paragraph(propre)
            if p.runs: p.runs[0].font.size = Pt(11)

# ============================================================
# EXÉCUTION
# ============================================================
print("\n🚀 DÉMARRAGE DU PROGRAMME 1 : ANALYSE GLOBALE & STRATÉGIQUE (GEMINI)")
texte_rc = extraire_texte_pdf("RC.pdf")
texte_ccap = extraire_texte_pdf("CCAP.pdf")
texte_cctp = extraire_texte_pdf("CCTP.pdf")

if not texte_cctp:
    print("❌ Le fichier CCTP.pdf est obligatoire.")
    exit()

# ÉTAPE 1 : FILTRAGE & NOTE DU JURY
print("\n📊 Étape 1/2 : Analyse d'éligibilité et Simulation Note du Jury...")
prompt_filtrage = f"""
Analyse ces documents de marché public et extrais les informations clés.
Calcule ensuite une simulation de la note théorique que le jury donnerait à N2IGC sur 100 points, basée sur une estimation de notre proposition (prix compétitifs, respect strict des délais, excellence technique).

Structure ta réponse EXACTEMENT ainsi :

DÉCISION: [ÉLIGIBLE ou REJETÉ]
RAISON: [Explication en 2 phrases]

## SIMULATION DE LA NOTE DU JURY (SUR 100)
- NOTE GLOBALE ESTIMÉE : [X/100]
- Note Critère Prix : [X/Points alloués] (Justification selon les exigences financières du RC)
- Note Critère Technique : [X/Points alloués] (Justification selon la complexité du CCTP)
- Note Critère Délais/RSE : [X/Points alloués]

## DONNÉES CLÉS EXTRAITES
- MONTANT_ESTIMÉ: [Montant ou "Non précisé"]
- VISITE_OBLIGATOIRE: [Oui / Non]
- DATE_LIMITE_OFFRES: [Date]
- DURÉE_MARCHÉ: [Durée]

--- RC ---
{texte_rc[:4000]}
--- CCAP ---
{texte_ccap[:4000]}
"""
resultat_filtrage = appeler_gemini(prompt_filtrage, "Tu es un expert en marchés publics français et analyste de risques.")

# ÉTAPE 2 : CAHIER DES CHARGES TECHNIQUE
print("\n🔧 Étape 2/2 : Génération du Cahier des Charges Technique...")
prompt_technique = f"""
Analyse ce CCTP et rédige un cahier des charges technique ultra-détaillé par lot/équipement pour N2IGC.

Pour CHAQUE lot identifié, respecte cette structure :
## [Nom du lot]
- Quantité requise : [Nombre]
- Contraintes fonctionnelles majeures : [Liste]
- Spécifications technologiques strictes : [Liste des dimensions, puissances, normes]
- Points de vigilance pour le sourcing : [Où est le piège ?]

--- CCTP ---
{texte_cctp[:12000]}
"""
resultat_technique = appeler_gemini(prompt_technique, "Tu es ingénieur principal spécialisé en systèmes didactiques et industriels.")

# CRÉATION DU RAPPORT WORD
print("\n📝 Génération des fichiers de sortie...")
doc = Document()
doc.add_heading("RAPPORT D'ANALYSE & CAHIER DES CHARGES - N2IGC", level=1)
doc.add_paragraph(f"Généré via Gemini le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}\n")

ajouter_titre(doc, "1. ÉLIGIBILITÉ ET PREDICTION NOTE JURY", niveau=1)
ecrire_lignes(doc, resultat_filtrage)
doc.add_page_break()

ajouter_titre(doc, "2. CAHIER DES CHARGES TECHNIQUE DÉTAILLÉ", niveau=1)
ecrire_lignes(doc, resultat_technique)

doc.save("N2IGC_Analyse_Marche.docx")
with open("N2IGC_Analyse_Marche.txt", "w", encoding="utf-8") as f:
    f.write("=== FILTRAGE & NOTE ===\n" + resultat_filtrage + "\n\n=== CCTP ===\n" + resultat_technique)

print("✅ PROGRAMME 1 TERMINÉ : Fichiers 'N2IGC_Analyse_Marche.docx' et '.txt' créés avec succès !")